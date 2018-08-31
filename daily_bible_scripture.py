#!/usr/bin/env python
# -*- coding: utf-8 -*-
import urllib,random
from lxml import etree
import datetime,time
import pandas as pd
import numpy as np
import calendar
import smtplib
from email.mime.text import MIMEText
from wxpy import *
import getpass
from access_google_sheet import from_google_sheet_to_txt
from docx import Document
from PIL import Image
from PIL import ImageFont
from PIL import ImageDraw

book_corr_lib={"路".decode("utf-8"):"Luke",
               "赛".decode("utf-8"):"Isaiah",
               "箴言".decode("utf-8"):"Proverbs",
               "帖前".decode("utf-8"):"1_Thessalonians",
               "诗篇".decode("utf-8"):"Psalm",
               "耶".decode("utf-8"):"Jeremiah",
               "约".decode("utf-8"):"John",
               "来".decode("utf-8"):"Hebrews",
               "哀".decode("utf-8"):"Lamentations",
               "结".decode("utf-8"):"Ezekiel",
               "雅".decode("utf-8"):"James",
               "彼前".decode("utf-8"):"1_Peter",
               "彼后".decode("utf-8"):"2_Peter",
               "约一".decode("utf-8"):"1_John",
               "约二".decode("utf-8"):"2_John",
               "约三".decode("utf-8"):"3_John",
               "犹".decode("utf-8"):"Jude",
               "何".decode("utf-8"):"Hosea",
               "珥".decode("utf-8"):"Joel",
               "摩".decode("utf-8"):"Amos",
               "俄".decode("utf-8"):"Obadiah",
               "启".decode("utf-8"):"Revelation",
               "拿".decode("utf-8"):"Jonah",
               "弥".decode("utf-8"):"Micah",
               "鸿".decode("utf-8"):"Nahum",
               "哈".decode("utf-8"):"Habakkuk",
               "番".decode("utf-8"):"Zephaniah",
               "该".decode("utf-8"):"Haggai",
               "亚".decode("utf-8"):"Zechariah",
               "玛".decode("utf-8"):"Malachi",
               "但".decode("utf-8"):"Daniel"}

#collection of (chaper,verse) of Book Proverbs to be shown underneath the message reminder each day!
chapter_verse_proverbs=[(1,7),(1,20),(1,33),(2,2),(2,6),(2,10),(2,20),(2,21),(3,5),(3,7),(3,13),(3,19),(3,27),(3,35),\
                        (4,8),(4,13),(8,12),(9,9),(9,10),(11,19),(11,25),(11,30),(12,25),(15,1),(15,4),(15,13),(15,23),\
                        (15,30),(15,30),(15,33),(16,9),(16,16),(16,20)]

def alignment(str1, space, align = 'left'):
    length = len(str1.encode('gb2312'))
    space = space - length if space >=length else 0
    if align == 'left':
        str1 = str1 + ' ' * space
    elif align == 'right':
        str1 = ' '* space +str1
    elif align == 'center':
        str1 = ' ' * (space //2) +str1 + ' '* (space - space // 2)
    return str1

send_wechat=raw_input("Send wechat reminder y or [n]:") or "n"
send_wechat = send_wechat=="y"
if send_wechat:
    bot=Bot()
key_today=None
wechat_friends=["群子在德国","全年读经运动"]
date_mode=raw_input("Use date specify mode [y] or n:") or "y"
if date_mode=="y":
    today_date=datetime.date.today() #today's date
    today_month,today_date=str(today_date.month),str(today_date.day)
    today_month=raw_input("Which month (1-12) or [today's month by default]:") or today_month
    today_date=raw_input("Which date (1-31) or [today's date by default]:") or today_date
    key_today="/".join([today_month,today_date])
else:
    today_date=datetime.date.today() #today's date
    today_month,today_date=str(today_date.month),str(today_date.day)
    book_chapter_verse=raw_input("Specify the book chapter and verses [赛,22:1-4]:") or "赛,22:1-4"
update_g_sheets=raw_input("Do you want to update bible reading plan from Google Sheet? y or [n]: ") or "n"
update_g_sheets = update_g_sheets=='y'
jason_credential_file="Worship-arrangement-DD-1005ad7eaf1f.json"#credential file for Google sheet API
#text file of worship service schedule.
bible_reading_plan='bible_reading_plan.txt'
#text file of Email corresponding of each person
if update_g_sheets:
    #now update these information from google spreadsheet
    print "Updating daily_bible_plan.txt from google sheet"
    try:
        from_google_sheet_to_txt(g_file_name="bible_reading_plan",save_file="bible_reading_plan.txt",sheet_tag="Sheet1",jason_credential_file=jason_credential_file)
        print "bible_reading_plan.txt is updated."
    except:
        print "Unable to download the google sheet for Sunday.You should manually download the sheet if it is different from your local copy!"
bible_reading = open(bible_reading_plan,'r').readlines()[1:]
bible_reading_lib={}
scripture_today={"oldtestament":"赛,22:1-4","newtestament":"路,22:1-4"}
for each in bible_reading:
    items=each.rstrip().rsplit()
    bible_reading_lib[items[0]]=items[1:]
if key_today!=None:
    scripture_today["oldtestament"]=bible_reading_lib[key_today][0]
    scripture_today["newtestament"]=bible_reading_lib[key_today][1]
else:
    scripture_today["oldtestament"]=book_chapter_verse
    scripture_today["newtestament"]="-"

cv_p=chapter_verse_proverbs[random.randint(0,31)]
sock=urllib.urlopen("http://www.chinesebibleonline.com/book/{0}/{1}".format("Proverbs",cv_p[0]))
htmlsource=sock.read()
sock.close()
s1=etree.HTML(htmlsource)
text_whole_chapter=s1.xpath('//*[@id="page_container"]/div[2]/div[position() mod2=0]/span/text()')
scripture_proverb=[each.encode('utf-8') for each in text_whole_chapter][cv_p[1]-1].rstrip().decode("utf8")

line1=alignment(("今天(%s月%s日)读经章节"%tuple([today_month,today_date])).decode("utf8"), 28, align = 'center')
line2=alignment(("-"*28).decode("utf8"), 28, align = 'center')
line3=alignment(("旧约：%s"%tuple([scripture_today["oldtestament"].replace(":*","")])).decode("utf8"), 28, align = 'center')
line4=alignment(("新约：%s"%tuple([scripture_today["newtestament"].replace(":*","")])).decode("utf8"), 28, align = 'center')
line5=alignment(("*"*28).decode("utf8"), 28, align = 'center')
line6=scripture_proverb
num_end=(len(line6)/2-1)
if len(line6)>=num_end:
    line6=alignment(line6[0:(len(line6)-num_end)], 28, align = 'center')+u"\n"+alignment((line6[(len(line6)-num_end):]), 28, align = 'center')
line7=alignment("+（：祝您读经快乐：）+".decode("utf8"), 28, align = 'center')
line8=alignment(("*"*28).decode("utf8"), 28, align = 'center')
TEXT="\n".join([line1,line2,line3,line4,line5,line6,line7,line8])
if send_wechat:
    print TEXT
    for friend in wechat_friends:
        temp_group=bot.search(friend.decode("utf8"))[0]
        temp_group.send_msg(TEXT)
        tag=random.randint(1,14)
        img = Image.open("base_images/base_img{0}.jpeg".format(tag))
        offset_lib={1:200,2:200,5:100,7:200,11:100,12:100,4:-400,9:-200,13:-200,14:-200,15:-100}
        if tag in offset_lib.keys():
            offset=offset_lib[tag]
        else:
            offset=0
        draw = ImageDraw.Draw(img)
        font = ImageFont.truetype("/Library/Fonts/Microsoft/SimSun.ttf", 23)
        draw.text((150, 650+offset),TEXT,(256,256,0),font=font)
        img.save('base_img_revised.jpg')
        temp_group.send_image("base_img_revised.jpg")
else:
    print TEXT
    offset_lib={1:200,2:200,5:100,7:200,11:100,12:100,4:-400,9:-200,13:-200,14:-200,15:-100}
    i=random.randint(1,14)
    img = Image.open("base_images/base_img{0}.jpeg".format(i))
    if i in offset_lib.keys():
        offset=offset_lib[i]
    else:
        offset=0
    draw = ImageDraw.Draw(img)
    font = ImageFont.truetype("/Library/Fonts/Microsoft/SimSun.ttf", 23)
    draw.text((150, 650+offset),TEXT,(256,256,0),font=font)
    img.save("base_img_revised.jpg")

temp_scripture_holder=[]
temp_scripture_holder.append("********************************%s月%s日读经章节***************************\n"%tuple([today_month,today_date]))
for bible_today_tag in scripture_today.values():
    if bible_today_tag=="-":
        pass
    else:
        print bible_today_tag
        book,chapter_verse=bible_today_tag.rsplit(",")
        chapters,verse=chapter_verse.rsplit(':')
        book=book_corr_lib[book.decode("utf-8")]
        chapters_boundary=map(int,chapters.rsplit('-'))
        if len(chapters_boundary)==2:
            chapters=range(chapters_boundary[0],chapters_boundary[1]+1)
        elif len(chapters_boundary)==1:
            chapters=range(chapters_boundary[0],chapters_boundary[0]+1)
        verse=verse.rsplit("-")
        if verse[0]!="*":
            verse=[int(verse[0]),int(verse[1])]
        for chapter in chapters:
            #print "Chapter {} of {}\n".format(chapter,book)
            temp_scripture_holder.append("Chapter {} of {}\n".format(chapter,book))
            sock=urllib.urlopen("http://www.chinesebibleonline.com/book/{0}/{1}".format(book,chapter))
            htmlsource=sock.read()
            sock.close()
            s1=etree.HTML(htmlsource)
            text_whole_chapter=s1.xpath('//*[@id="page_container"]/div[2]/div[position() mod2=0]/span/text()')
            text_whole_chapter=[each.encode('utf-8').decode("utf-8") for each in text_whole_chapter]
            if verse[0]=="*":
                for i in range(len(text_whole_chapter)):
                    #print i+1,text_whole_chapter[i]
                    temp_scripture_holder.append('{0}.{1}'.format(i+1,text_whole_chapter[i].encode("utf-8")))
            else:
                for i in range(verse[0]-1,verse[1]):
                    #print i+1,text_whole_chapter[i]
                    temp_scripture_holder.append("{0}.{1}".format(i+1,text_whole_chapter[i].encode("utf-8")))

document_today = Document()
document_today.add_paragraph("".join(temp_scripture_holder).decode("utf8"))
document_today.save('current scripture.docx')
if today_month==str(datetime.date.today().month) and today_date==str(datetime.date.today().day) and date_mode=="y":
    accumulated_dates=[]
    with open("accumulated_date.txt","r") as write_f:
        for each in write_f.readlines():
            accumulated_dates.append(each.rstrip())
    if key_today not in accumulated_dates:
        document_today = Document()
        document_today.add_paragraph("".join(temp_scripture_holder).decode("utf8"))
        document_today.save('today scripture.docx')
        document_acc = Document('scriptures accumulated.docx')
        document_acc.add_paragraph("".join(temp_scripture_holder).decode("utf8"))
        document_acc.save('scriptures accumulated.docx')
        accumulated_dates.append(key_today)
        with open("accumulated_date.txt","w") as write_f:
            write_f.write("\n".join(accumulated_dates))
